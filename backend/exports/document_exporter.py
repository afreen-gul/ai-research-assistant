"""Document export to PDF, DOCX, and Markdown."""

import io
import logging
import re
from typing import Any

import markdown as md
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

logger = logging.getLogger(__name__)


def _strip_citation_markers_for_plain(text: str) -> str:
    return re.sub(r"\[(\d+)\]", r"[\1]", text)


def export_markdown(content: str, title: str = "Document", bibliography: list[str] | None = None) -> str:
    output = f"# {title}\n\n{content}\n"
    if bibliography:
        output += "\n\n## References\n\n"
        for entry in bibliography:
            output += f"- {entry}\n"
    return output


def export_docx(content: str, title: str = "Document", bibliography: list[str] | None = None) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, 0)

    for para in content.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith("# "):
            doc.add_heading(para[2:], level=1)
        elif para.startswith("## "):
            doc.add_heading(para[3:], level=2)
        elif para.startswith("### "):
            doc.add_heading(para[4:], level=3)
        elif para.startswith("|") and "|" in para[1:]:
            rows = [r.strip() for r in para.split("\n") if r.strip()]
            if len(rows) >= 2:
                headers = [c.strip() for c in rows[0].strip("|").split("|")]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for row in rows[2:]:
                    cells = [c.strip() for c in row.strip("|").split("|")]
                    row_cells = table.add_row().cells
                    for i, c in enumerate(cells[:len(headers)]):
                        row_cells[i].text = c
            else:
                doc.add_paragraph(para)
        else:
            doc.add_paragraph(para)

    if bibliography:
        doc.add_heading("References", level=1)
        for entry in bibliography:
            p = doc.add_paragraph(entry)
            p.style.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(content: str, title: str = "Document", bibliography: list[str] | None = None) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CustomTitle", parent=styles["Heading1"], fontSize=18, spaceAfter=20)
    body_style = ParagraphStyle("CustomBody", parent=styles["Normal"], fontSize=11, leading=14, spaceAfter=10)
    ref_style = ParagraphStyle("Ref", parent=styles["Normal"], fontSize=9, leading=12, spaceAfter=6)

    story: list[Any] = [Paragraph(title, title_style), Spacer(1, 12)]

    for para in content.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if para.startswith("# "):
            story.append(Paragraph(para[2:], styles["Heading1"]))
        elif para.startswith("## "):
            story.append(Paragraph(para[3:], styles["Heading2"]))
        elif para.startswith("### "):
            story.append(Paragraph(para[4:], styles["Heading3"]))
        else:
            story.append(Paragraph(para, body_style))
        story.append(Spacer(1, 6))

    if bibliography:
        story.append(Spacer(1, 20))
        story.append(Paragraph("References", styles["Heading2"]))
        for entry in bibliography:
            story.append(Paragraph(entry, ref_style))

    doc.build(story)
    return buffer.getvalue()
